from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from .utilities import strip_md

def field(paragraph, instr):
    run=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), instr); run._r.append(fld)

def add_header_footer(sec):
    p=sec.header.paragraphs[0]; p.text='Take Control of Your Inbox    |    Outlook on iPadOS'; p.style='Header'
    tbl=sec.footer.add_table(rows=1, cols=3, width=Inches(7)); tbl.style='Table Grid'
    vals=['Take Control of Your Inbox','Outlook on iPadOS','Page ']
    for i,v in enumerate(vals):
        para=tbl.cell(0,i).paragraphs[0]; para.text=v
        if i==2: field(para,'PAGE')

def add_toc(doc):
    p=doc.add_paragraph(); p.add_run('Table of contents').bold=True; p.style='Heading 1'
    field(doc.add_paragraph(), 'TOC \\o "1-2" \\h \\z \\u')

def placeholder(doc, text):
    t=doc.add_table(rows=1, cols=1); t.style='Table Grid'; c=t.cell(0,0); c.text=text
    for p in c.paragraphs:
        for r in p.runs: r.font.size=Pt(9.5)

def add_md_block(doc, block):
    lines=block.splitlines(); in_table=[]
    def flush_table():
        nonlocal in_table
        if len(in_table)>1:
            rows=[]
            for l in in_table:
                if re.match(r'^\s*\|?\s*:?-{3,}',l): continue
                cells=[strip_md(x) for x in l.strip().strip('|').split('|')]
                rows.append(cells)
            if rows:
                tbl=doc.add_table(rows=len(rows), cols=max(map(len,rows))); tbl.style='Table Grid'
                for r,row in enumerate(rows):
                    for c,val in enumerate(row): tbl.cell(r,c).text=val
        in_table=[]
    for raw in lines:
        line=raw.rstrip()
        if '|' in line and line.strip().startswith('|'):
            in_table.append(line); continue
        flush_table()
        if not line.strip(): continue
        m=re.match(r'^(#{1,6})\s+(.*)',line)
        if m:
            txt=strip_md(m.group(2))
            if re.match(r'Page \d+',txt):
                doc.add_paragraph(txt, style='Heading 2')
            elif len(m.group(1))<=2: doc.add_paragraph(txt, style='Heading 1')
            else: doc.add_paragraph(txt, style='Heading 3')
        elif re.match(r'^\s*[-*]\s+',line): doc.add_paragraph(strip_md(re.sub(r'^\s*[-*]\s+','',line)), style='List Bullet')
        elif re.match(r'^\s*\d+[.)]\s+',line): doc.add_paragraph(strip_md(re.sub(r'^\s*\d+[.)]\s+','',line)), style='List Number')
        elif re.search(r'(Screenshot|visual|image) needed|Suggested visual|Design notes', line, re.I):
            placeholder(doc, 'Screenshot needed / production note: '+strip_md(line)+'\nAlt-text draft: To be finalized during screenshot capture.\nValidation note: Use only approved repository screenshots.')
        else: doc.add_paragraph(strip_md(line), style='Normal')
    flush_table()
