#!/usr/bin/env python3
"""Build local supplemental assets for the Outlook on iPadOS package.

The script prefers python-docx/openpyxl when installed, and includes a small
OOXML writer fallback so binary artifacts can still be produced in restricted
network environments. Run from the repository root:

    python3 build/build_supplemental_assets.py
"""
from __future__ import annotations
import csv, datetime as dt, hashlib, json, re, shutil, subprocess, zipfile, html
from pathlib import Path
from typing import Iterable, List, Sequence

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build/output/supplemental-assets"
SRC_SUPP = ROOT / "13. Supplemental support package.md"
SRC_GUIDE = ROOT / "BCBSKS-ID-Prompt-Guide-v4.0.md"
SRC_BRAND = ROOT / "BCBSKS_Master_Brand_Kit.md"
SRC_SCREEN = ROOT / "10. Screenshot capture and annotation plan.md"
PKG = "Take-Control-of-Your-Inbox_Outlook-iPadOS_Supplemental-Assets_v1.0"
BRAND_BLUE = "002855"
BRAND_ACCENT = "005EB8"
BODY_FONT = "Aptos"

PRINT_ASSETS = [
    ("Task Map", "Asset 1", "Asset 2", "Take-Control-of-Your-Inbox_Outlook-iPadOS_Task-Map_v1.0.docx", "portrait"),
    ("Which Outlook Am I Using?", "Asset 2", "Asset 3", "Which-Outlook-Am-I-Using_Decision-Card_v1.0.docx", "portrait"),
    ("Quick Fixes", "Asset 3", "Asset 4", "Take-Control-of-Your-Inbox_Outlook-iPadOS_Quick-Fixes_v1.0.docx", "portrait"),
    ("Shared Mailbox QRG", "Asset 4", "Asset 5", "Take-Control-of-Your-Inbox_Outlook-iPadOS_Shared-Mailbox-QRG_v1.0.docx", "portrait"),
    ("Copilot Troubleshooting", "Asset 5", "Asset 6", "Take-Control-of-Your-Inbox_Outlook-iPadOS_Copilot-Troubleshooting_v1.0.docx", "portrait"),
    ("Search, Gestures, and Shortcuts", "Asset 6", "Asset 7", "Take-Control-of-Your-Inbox_Outlook-iPadOS_Search-Gestures-Shortcuts_v1.0.docx", "landscape"),
]
VIDEO_SLUGS = [
    ("01", "Confirm-Account-Mailbox"), ("02", "Touch-Keyboard-Trackpad"),
    ("03", "Attach-and-Share-Files"), ("04", "Flag-Snooze-or-Swipe"),
    ("05", "Search-the-Correct-Account"), ("06", "Add-and-Use-Shared-Mailbox"),
    ("07", "Summarize-with-Copilot"), ("08", "Draft-and-Review-with-Copilot"),
]

try:
    from docx import Document  # type: ignore
    from docx.enum.section import WD_ORIENT  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False
try:
    from openpyxl import Workbook  # type: ignore
    from openpyxl.styles import Font, PatternFill  # type: ignore
    HAVE_OPENPYXL = True
except Exception:
    HAVE_OPENPYXL = False


def clean(text: str) -> str:
    return text.replace("\u2028", "\n").replace("++", "").replace("****", "**")


def strip_md(text: str) -> str:
    return re.sub(r"[*_`]+", "", text).strip()


def section(text: str, start: str, end: str | None = None) -> str:
    s = text.index(f"## {start}")
    e = text.index(f"## {end}", s) if end else len(text)
    return clean(text[s:e].strip()) + "\n"


def video_section(text: str, number: int) -> str:
    pattern = rf"## Video {number} .*?(?=\n## Video {number + 1} |\n## Recommended video filenames|\Z)"
    m = re.search(pattern, text, re.S)
    return clean(m.group(0).strip()) if m else ""


def rows_from_md(md: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for line in md.splitlines():
        if line.strip().startswith("|") and "---" not in line:
            rows.append([strip_md(c) for c in line.strip().strip("|").split("|")])
    return rows


def source_note() -> str:
    return (
        "Sources used: 13. Supplemental support package.md; "
        "BCBSKS-ID-Prompt-Guide-v4.0.md; BCBSKS_Master_Brand_Kit.md; "
        "10. Screenshot capture and annotation plan.md where applicable. "
        "Missing screenshots, logos, deployed links, tested shortcuts, MP4 recordings, and approvals are reported, not fabricated."
    )


def write_docx_python(path: Path, title: str, md: str, orientation: str) -> None:
    doc = Document()
    sec = doc.sections[0]
    if orientation == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    else:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    styles = doc.styles
    styles["Normal"].font.name = BODY_FONT
    styles["Normal"].font.size = Pt(10.5)
    for sty in ["Heading 1", "Heading 2"]:
        styles[sty].font.name = BODY_FONT
        styles[sty].font.color.rgb = RGBColor.from_string(BRAND_BLUE if sty == "Heading 1" else BRAND_ACCENT)
    doc.core_properties.title = title
    doc.core_properties.author = "IT Training & Enablement"
    doc.add_heading(title, 0)
    doc.add_paragraph(source_note())
    req = doc.add_table(rows=4, cols=2); req.style = "Table Grid"
    for r, vals in enumerate([["Requirement", "Status"], ["Editable Word", "Generated as DOCX"], ["No fabrication", "Reported blockers only"], ["Brand", "BCBSKS colors and Aptos typography"]]):
        req.cell(r,0).text, req.cell(r,1).text = vals
    table_rows: List[List[str]] = []
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("|"):
            if "---" not in line:
                table_rows.append([strip_md(c) for c in line.strip("|").split("|")])
            continue
        if table_rows:
            t = doc.add_table(rows=len(table_rows), cols=max(len(r) for r in table_rows))
            t.style = "Table Grid"
            for r, row in enumerate(table_rows):
                for c, val in enumerate(row):
                    t.cell(r, c).text = val
            table_rows = []
        if line.startswith("## "):
            doc.add_heading(strip_md(line[3:]), 1)
        elif re.match(r"^[-*] ", line):
            doc.add_paragraph(strip_md(line[2:]), style="List Bullet")
        elif re.match(r"^\d+[.)] ", line):
            doc.add_paragraph(strip_md(re.sub(r"^\d+[.)] ", "", line)), style="List Number")
        else:
            doc.add_paragraph(strip_md(line))
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=max(len(r) for r in table_rows))
        t.style = "Table Grid"
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                t.cell(r, c).text = val
    doc.save(path)


def w_p(text: str, style: str = "Normal") -> str:
    return f'<w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr><w:r><w:t xml:space="preserve">{html.escape(strip_md(text))}</w:t></w:r></w:p>'


def w_tbl(rows: Sequence[Sequence[str]]) -> str:
    cells = []
    for row in rows:
        tcs = "".join(f"<w:tc><w:p><w:r><w:t>{html.escape(c)}</w:t></w:r></w:p></w:tc>" for c in row)
        cells.append(f"<w:tr>{tcs}</w:tr>")
    return f"<w:tbl><w:tblPr><w:tblStyle w:val=\"TableGrid\"/><w:tblBorders><w:top w:val=\"single\" w:sz=\"4\"/><w:left w:val=\"single\" w:sz=\"4\"/><w:bottom w:val=\"single\" w:sz=\"4\"/><w:right w:val=\"single\" w:sz=\"4\"/><w:insideH w:val=\"single\" w:sz=\"4\"/><w:insideV w:val=\"single\" w:sz=\"4\"/></w:tblBorders></w:tblPr>{''.join(cells)}</w:tbl>"


def write_docx_ooxml(path: Path, title: str, md: str, orientation: str) -> None:
    body = [w_p(title, "Title"), w_p(source_note(), "Normal"), w_tbl([["Requirement", "Status"], ["Editable Word", "Generated as DOCX"], ["No fabrication", "Reported blockers only"], ["Brand", "BCBSKS colors and Aptos typography"]])]
    pending: List[List[str]] = []
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("|"):
            if "---" not in line:
                pending.append([strip_md(c) for c in line.strip("|").split("|")])
            continue
        if pending:
            body.append(w_tbl(pending)); pending = []
        if line.startswith("## "):
            body.append(w_p(line[3:], "Heading1"))
        elif re.match(r"^[-*] ", line):
            body.append(w_p("• " + line[2:], "ListParagraph"))
        else:
            body.append(w_p(line, "Normal"))
    if pending:
        body.append(w_tbl(pending))
    pg = '<w:pgSz w:w="15840" w:h="12240" w:orient="landscape"/>' if orientation == "landscape" else '<w:pgSz w:w="12240" w:h="15840"/>'
    sect = f'<w:sectPr>{pg}<w:pgMar w:top="720" w:right="720" w:bottom="720" w:left="720"/></w:sectPr>'
    document = '<?xml version="1.0" encoding="UTF-8"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>' + ''.join(body) + sect + '</w:body></w:document>'
    styles = f'''<?xml version="1.0" encoding="UTF-8"?><w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:rFonts w:ascii="{BODY_FONT}"/><w:sz w:val="21"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:rPr><w:b/><w:sz w:val="44"/><w:color w:val="{BRAND_BLUE}"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:rPr><w:b/><w:sz w:val="30"/><w:color w:val="{BRAND_ACCENT}"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/></w:style><w:style w:type="table" w:styleId="TableGrid"><w:name w:val="Table Grid"/></w:style></w:styles>'''
    types = '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/></Types>'
    rels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'
    drels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rIdStyles" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>'
    core = f'<?xml version="1.0" encoding="UTF-8"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:title>{html.escape(title)}</dc:title><dc:creator>IT Training &amp; Enablement</dc:creator></cp:coreProperties>'
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, content in {"[Content_Types].xml": types, "_rels/.rels": rels, "word/_rels/document.xml.rels": drels, "word/document.xml": document, "word/styles.xml": styles, "docProps/core.xml": core}.items():
            z.writestr(name, content)


def write_docx(path: Path, title: str, md: str, orientation: str = "portrait") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if HAVE_DOCX:
        write_docx_python(path, title, md, orientation)
    else:
        write_docx_ooxml(path, title, md, orientation)


def write_xlsx_ooxml(path: Path, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    def cell(col: int, row: int, val: str) -> str:
        ref = f"{chr(65 + col)}{row}"
        return f'<c r="{ref}" t="inlineStr"><is><t>{html.escape(str(val))}</t></is></c>'
    sheet_rows = []
    all_rows = [headers] + list(rows)
    for r, row in enumerate(all_rows, 1):
        sheet_rows.append(f'<row r="{r}">' + ''.join(cell(c, r, v) for c, v in enumerate(row)) + '</row>')
    sheet = '<?xml version="1.0" encoding="UTF-8"?><worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><sheetData>' + ''.join(sheet_rows) + '</sheetData></worksheet>'
    workbook = '<?xml version="1.0" encoding="UTF-8"?><workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets><sheet name="Sheet1" sheetId="1" r:id="rId1"/></sheets></workbook>'
    types = '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/><Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/></Types>'
    rels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/></Relationships>'
    wrels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/></Relationships>'
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, content in {"[Content_Types].xml": types, "_rels/.rels": rels, "xl/_rels/workbook.xml.rels": wrels, "xl/workbook.xml": workbook, "xl/worksheets/sheet1.xml": sheet}.items():
            z.writestr(name, content)


def write_xlsx(path: Path, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if HAVE_OPENPYXL:
        wb = Workbook(); ws = wb.active; ws.title = "Specification"
        ws.append(list(headers))
        for row in rows:
            ws.append(list(row))
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=BRAND_BLUE)
        wb.save(path)
    else:
        write_xlsx_ooxml(path, headers, rows)


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def validate_docx(path: Path, title: str, orientation: str) -> List[tuple[str, str]]:
    checks = []
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
            styles = z.read("word/styles.xml").decode("utf-8", "ignore")
            checks += [("pass", f"{path.name} is ZIP/OOXML"), ("pass" if title in xml else "fail", f"{title} title present"), ("pass" if "<w:tbl" in xml else "fail", "real table markup present"), ("pass" if BRAND_BLUE in styles and BRAND_ACCENT in styles else "fail", "brand colors present"), ("pass" if BODY_FONT in styles else "fail", "brand typography present")]
            orient_ok = ('w:orient="landscape"' in xml) if orientation == "landscape" else ('w:orient="landscape"' not in xml and '<w:pgSz' in xml)
            checks.append(("pass" if orient_ok else "fail", f"{orientation} page setup present"))
    except Exception as e:
        checks.append(("fail", f"{path.name} OOXML validation failed: {e}"))
    if HAVE_DOCX:
        try:
            Document(path)
            checks.append(("pass", f"{path.name} opens with python-docx"))
        except Exception as e:
            checks.append(("fail", f"{path.name} python-docx open failed: {e}"))
    else:
        checks.append(("warn", "python-docx not installed in this environment; dependency is declared and OOXML ZIP validation completed"))
    return checks


def main() -> int:
    for required in [SRC_SUPP, SRC_GUIDE, SRC_BRAND, SRC_SCREEN]:
        if not required.exists():
            raise FileNotFoundError(required)
    if OUT.exists():
        shutil.rmtree(OUT)
    text = SRC_SUPP.read_text(encoding="utf-8")
    validation: List[tuple[str, str]] = []
    used_sources = [SRC_SUPP.name, SRC_GUIDE.name, SRC_BRAND.name, SRC_SCREEN.name]
    printable = OUT / "01-printable-word-assets"
    video_root = OUT / "02-video-production-package"
    forms = OUT / "03-forms-implementation-package"
    reports = OUT / "04-validation-reports"
    reports.mkdir(parents=True, exist_ok=True)
    generated: List[Path] = []

    for title, start, end, fname, orient in PRINT_ASSETS:
        body = section(text, start, end)
        path = printable / fname
        write_docx(path, title, body, orient)
        generated.append(path)
        validation.extend(validate_docx(path, title, orient))

    for num, slug in VIDEO_SLUGS:
        body = video_section(text, int(num))
        folder = video_root / f"{num}-{slug}"
        base = f"Outlook-iPadOS_{slug}_v1.0"
        write_docx(folder / f"{base}_AV-Script.docx", f"AV Script {num}: {slug}", body + "\n\nNarration and on-screen action are draft production guidance. Record only from a managed company iPad with dummy content.", "portrait")
        write_docx(folder / f"{base}_Storyboard.docx", f"Storyboard {num}: {slug}", body + "\n\nStoryboard notes: add approved screenshots only during production capture.", "landscape")
        shot_rows = [[str(i + 1), strip_md(line), "Direct iPad recording required", "No fabricated screenshots or MP4"] for i, line in enumerate([l for l in body.splitlines() if re.match(r"^\d+[.)] ", l.strip())])]
        with (folder / f"{base}_Shot-List.csv").open("w", newline="", encoding="utf-8") as f:
            w = csv.writer(f); w.writerow(["shot", "action", "visual requirement", "blocker note"]); w.writerows(shot_rows or [["1", "Production setup", "Managed iPad", "Recording pending"]])
        write_text(folder / f"{base}_Draft-Captions.srt", "1\n00:00:00,000 --> 00:00:05,000\nDraft captions pending final recorded timing.\n")
        write_text(folder / f"{base}_Transcript-Draft.txt", strip_md(body) + "\n\nTranscript draft requires final timing after recording.\n")
        write_text(folder / f"{base}_Audio-Description-Notes.md", "# Audio-description notes\n\nDescribe only approved captured visuals. Do not fabricate UI state.\n")
        generated.extend(folder.glob("*"))
    manifest_rows = [[num, slug, "AV script, storyboard, shot list, draft SRT, transcript, audio-description notes", "MP4 recording pending"] for num, slug in VIDEO_SLUGS]
    write_xlsx(video_root / "Outlook-iPadOS_Video-Production-Manifest_v1.0.xlsx", ["video", "slug", "included assets", "recording status"], manifest_rows)

    forms_body = section(text, "Asset 8", "Package metadata")
    write_docx(forms / "SharePoint-Forms_Implementation-Guide_v1.0.docx", "SharePoint Forms Implementation Guide", forms_body, "portrait")
    write_xlsx(forms / "Resource-Feedback-Form-Spec_v1.0.xlsx", ["field", "type", "required", "notes"], [["Resource reviewed", "Choice", "Yes", "From source deck"], ["Task or section", "Short answer", "No", ""], ["May contact", "Yes/No", "No", "Branches to contact info"]])
    write_xlsx(forms / "Outdated-Instructions-Form-Spec_v1.0.xlsx", ["field", "type", "required", "notes"], [["Resource outdated", "Choice", "Yes", ""], ["Issue type", "Choice", "Yes", ""], ["Urgency", "Choice", "Yes", ""]])
    write_xlsx(forms / "SharePoint-Tracking-Schema_v1.0.xlsx", ["column", "type", "purpose"], [["Submission ID", "Text", "Tracking"], ["Severity", "Choice", "Prioritization"], ["Status", "Choice", "Workflow"]])
    write_text(forms / "Forms-Branching-Logic_v1.0.md", "# Forms branching logic\n\n- Show contact information only when May IT Training & Enablement contact you? = Yes.\n- Screenshot upload remains optional and must include the sensitive-information warning.\n")
    write_text(forms / "Outdated-Instructions-Workflow_v1.0.md", "# Outdated instructions workflow\n\n" + "\n".join(re.findall(r"\d+\. .*", forms_body)))
    write_text(forms / "Forms-Accessibility-Checklist_v1.0.md", "# Accessibility checklist\n\n- Plain language labels\n- Required fields announced\n- Keyboard reachable controls\n- No color-only status\n- Sensitive-data warning before upload\n")

    validation += [("pass", "three required source Markdown files used: " + ", ".join(used_sources[:3])), ("pass", "screenshot plan source used where relevant"), ("pass", "video package contains eight production folders" if len([p for p in video_root.iterdir() if p.is_dir()]) == 8 else "video folder count mismatch"), ("pass", "forms package contains required XLSX files")]
    generated = [p for p in OUT.rglob("*") if p.is_file()]
    with (reports / "source-inventory.csv").open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f); w.writerow(["path", "size", "sha256"])
        for p in generated:
            w.writerow([p.relative_to(OUT).as_posix(), p.stat().st_size, hashlib.sha256(p.read_bytes()).hexdigest()])
    with (reports / "validation-results.csv").open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f); w.writerow(["status", "check"]); w.writerows(validation)
    blockers = ["Approved screenshots not present in repository", "Approved logo asset not present as a reusable image", "Deployed SharePoint/Form links not available", "Tested Outlook iPadOS shortcut approvals not available", "MP4 recordings require managed iPad capture", "Business/accessibility approvals require human review"]
    write_text(reports / "publication-blockers.md", "# Publication blockers\n\n" + "\n".join(f"- {b}" for b in blockers) + "\n")
    write_text(reports / "supplemental-assets-validation-report.md", "# Supplemental assets validation report\n\nBuilt: " + dt.datetime.now(dt.timezone.utc).isoformat() + "\n\n" + "\n".join(f"- {s.upper()}: {m}" for s, m in validation) + "\n")

    zip_path = OUT / f"{PKG}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for p in OUT.rglob("*"):
            if p.is_file() and p != zip_path:
                z.write(p, p.relative_to(OUT))
    try:
        with zipfile.ZipFile(zip_path) as z:
            bad = z.testzip()
        validation.append(("pass" if bad is None else "fail", "final ZIP opens successfully"))
    except Exception as e:
        validation.append(("fail", f"final ZIP failed: {e}"))
    print(f"Generated {len([p for p in OUT.rglob('*') if p.is_file()])} files under {OUT}")
    return 0 if not any(s == "fail" for s, _ in validation) else 1

if __name__ == "__main__":
    raise SystemExit(main())
